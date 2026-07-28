#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown → Word(.docx),视觉与工作台的 md 预览保持一致。

为什么不用 textutil:它把 HTML <table> 拆成普通段落(实测四种写法均不支持),
而售前文档(简报/测算/评审)几乎都靠表格,所以改用 python-docx 直接生成。

用法:python3 bin/md2docx.py <input.md> <output.docx>
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT_CN = "PingFang SC"
FONT_MONO = "Menlo"
INK = RGBColor(0x16, 0x18, 0x1D)
MUTED = RGBColor(0x5B, 0x62, 0x70)
ACCENT = RGBColor(0x1F, 0x3D, 0x99)
H_SIZE = {1: 17, 2: 14, 3: 12.5, 4: 11.5, 5: 11, 6: 11}


def _set_font(run, name=FONT_CN, size=11, bold=False, italic=False,
              color=INK, mono=False):
    f = run.font
    f.name = FONT_MONO if mono else name
    f.size = Pt(size)
    f.bold = bold
    f.italic = italic
    f.color.rgb = color
    # 中文字形必须单独指定 eastAsia,否则 Word 用默认宋体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO if mono else name)


def _shade(cell, hex_color):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(el)


def _para_border(p, side="left", color="1F3D99", sz=18):
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    e = OxmlElement(f"w:{side}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "8")
    e.set(qn("w:color"), color)
    bd.append(e)
    pPr.append(bd)


TOKEN = re.compile(r"(\*\*[^*]+\*\*|(?<![\w*])\*[^*\n]+\*(?![\w*])|`[^`]+`|"
                   r"\[[^\]]+\]\(https?://[^)]+\))")


def add_rich(p, text, size=11, base_bold=False, color=INK):
    """把一行 Markdown 内联标记渲染成多个 run(粗体/斜体/行内代码/链接)。"""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _set_font(p.add_run(part[2:-2]), size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            _set_font(r, size=size - 1, mono=True, color=color)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            _set_font(p.add_run(part[1:-1]), size=size, italic=True,
                      bold=base_bold, color=color)
        elif part.startswith("[") and "](" in part:
            label, url = part[1:-1].split("](", 1)
            _set_font(p.add_run(label), size=size, bold=base_bold, color=ACCENT)
        else:
            _set_font(p.add_run(part), size=size, bold=base_bold, color=color)


def convert(md_path, out_path):
    md = Path(md_path).read_text(encoding="utf-8", errors="replace")
    # 去掉 frontmatter
    md = re.sub(r"^---\s*\n.*?\n---\s*\n", "", md, count=1, flags=re.S)
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT_CN
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.5

    lines = md.split("\n")
    i, in_code, code_buf = 0, False, []
    while i < len(lines):
        L = lines[i]

        if in_code:
            if L.startswith("```"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(10)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                _set_font(p.add_run("\n".join(code_buf)), size=9.5, mono=True, color=MUTED)
                in_code, code_buf = False, []
            else:
                code_buf.append(L)
            i += 1
            continue
        if L.startswith("```"):
            in_code = True
            i += 1
            continue

        if not L.strip():
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)", L)
        if m:
            lv = len(m.group(1))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12 if lv <= 2 else 9)
            p.paragraph_format.space_after = Pt(4)
            add_rich(p, m.group(2), size=H_SIZE.get(lv, 11), base_bold=True)
            for r in p.runs:
                r.font.bold = True
            i += 1
            continue

        # 分隔线
        if re.match(r"^([-*_]\s*){3,}$", L):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            _para_border(p, "bottom", "D5D9E0", 6)
            i += 1
            continue

        # 引用(连续多行合并)
        if L.startswith(">"):
            buf = [L.lstrip("> ")]
            while i + 1 < len(lines) and lines[i + 1].startswith(">"):
                i += 1
                buf.append(lines[i].lstrip("> "))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.space_before = Pt(4)
            _para_border(p, "left", "1F3D99", 18)
            add_rich(p, " ".join(buf), size=10.5, color=MUTED)
            i += 1
            continue

        # 表格
        if "|" in L and i + 1 < len(lines) and \
                re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", lines[i + 1]):
            def cells(row):
                return [c.strip() for c in row.strip().strip("|").split("|")]
            header = cells(L)
            body = []
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                body.append(cells(lines[i]))
                i += 1
            ncol = len(header)
            t = doc.add_table(rows=1, cols=ncol)
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            for j, h in enumerate(header):
                c = t.rows[0].cells[j]
                c.text = ""
                _shade(c, "EEF0F3")
                add_rich(c.paragraphs[0], h, size=10, base_bold=True)
                for r in c.paragraphs[0].runs:
                    r.font.bold = True
            for row in body:
                cs = t.add_row().cells
                for j in range(ncol):
                    cs[j].text = ""
                    add_rich(cs[j].paragraphs[0], row[j] if j < len(row) else "", size=10)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # 列表(支持缩进与任务框)
        mu = re.match(r"^(\s*)[-*+]\s+(.*)", L)
        mo = re.match(r"^(\s*)\d+[.、]\s+(.*)", L)
        if mu or mo:
            g = mu or mo
            depth = min(len(g.group(1)) // 2, 3)
            body = g.group(2)
            body = re.sub(r"^\[\s\]\s*", "☐ ", body)
            body = re.sub(r"^\[x\]\s*", "☑ ", body, flags=re.I)
            style = "List Bullet" if mu else "List Number"
            try:
                p = doc.add_paragraph(style=style if depth == 0 else f"{style} {depth+1}")
            except KeyError:
                p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Pt(18 + depth * 14)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.4
            add_rich(p, body, size=11)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_rich(p, L, size=11)
        i += 1

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python3 bin/md2docx.py <input.md> <output.docx>")
        sys.exit(1)
    print(convert(sys.argv[1], sys.argv[2]))
